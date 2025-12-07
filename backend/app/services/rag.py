"""
RAG (Retrieval-Augmented Generation) Service
Handles document embedding and semantic search for policy documents
"""

from typing import List, Optional, Dict, Any
import structlog

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, text
from pgvector.sqlalchemy import Vector

from app.config import settings
from app.models import PolicyDocument

logger = structlog.get_logger()

# Lazy load sentence transformers to avoid startup overhead
_embedding_model = None


def get_embedding_model():
    """Lazy load the embedding model."""
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer
        _embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL)
        logger.info("Loaded embedding model", model=settings.EMBEDDING_MODEL)
    return _embedding_model


class RAGService:
    """Service for RAG operations on policy documents."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
    
    async def embed_text(self, text: str) -> List[float]:
        """Generate embedding for a text string."""
        model = get_embedding_model()
        embedding = model.encode(text, convert_to_numpy=True)
        return embedding.tolist()
    
    async def search(
        self, 
        query: str, 
        policy_type: Optional[str] = None,
        top_k: int = None
    ) -> List[Dict[str, Any]]:
        """
        Search policy documents using semantic similarity.
        
        Args:
            query: Search query
            policy_type: Optional filter by policy type
            top_k: Number of results to return
            
        Returns:
            List of matching documents with content and metadata
        """
        top_k = top_k or settings.RAG_TOP_K_RESULTS
        
        try:
            # Generate query embedding
            query_embedding = await self.embed_text(query)
            
            # Build the search query using pgvector's cosine distance
            # Using raw SQL for vector operations
            sql = text("""
                SELECT 
                    id,
                    document_name,
                    document_type,
                    policy_type,
                    content,
                    chunk_index,
                    metadata,
                    1 - (embedding <=> :query_embedding::vector) as similarity
                FROM policy_documents
                WHERE (:policy_type IS NULL OR policy_type = :policy_type)
                ORDER BY embedding <=> :query_embedding::vector
                LIMIT :top_k
            """)
            
            result = await self.db.execute(
                sql,
                {
                    "query_embedding": str(query_embedding),
                    "policy_type": policy_type,
                    "top_k": top_k
                }
            )
            
            rows = result.fetchall()
            
            return [
                {
                    "id": str(row.id),
                    "document_name": row.document_name,
                    "document_type": row.document_type,
                    "policy_type": row.policy_type,
                    "content": row.content,
                    "chunk_index": row.chunk_index,
                    "relevance_score": float(row.similarity),
                    "metadata": row.metadata or {}
                }
                for row in rows
                if row.similarity > 0.3  # Minimum relevance threshold
            ]
            
        except Exception as e:
            logger.error("RAG search error", error=str(e))
            return []
    
    async def add_document(
        self,
        document_name: str,
        document_type: str,
        content: str,
        policy_type: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[str]:
        """
        Add a document to the RAG system.
        Chunks the document and creates embeddings.
        
        Args:
            document_name: Name of the document
            document_type: Type (pdf, docx, txt)
            content: Document content
            policy_type: Type of policy this document relates to
            metadata: Additional metadata
            
        Returns:
            List of created document chunk IDs
        """
        # Chunk the content
        chunks = self._chunk_text(content)
        created_ids = []
        
        for idx, chunk in enumerate(chunks):
            # Generate embedding
            embedding = await self.embed_text(chunk)
            
            # Create document record
            doc = PolicyDocument(
                document_name=document_name,
                document_type=document_type,
                policy_type=policy_type,
                chunk_index=idx,
                content=chunk,
                embedding=embedding,
                metadata=metadata or {}
            )
            
            self.db.add(doc)
            created_ids.append(str(doc.id))
        
        await self.db.commit()
        
        logger.info(
            "Document added to RAG", 
            document_name=document_name,
            chunks=len(chunks)
        )
        
        return created_ids
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks with overlap.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        chunk_size = settings.RAG_CHUNK_SIZE
        overlap = settings.RAG_CHUNK_OVERLAP
        
        # Simple chunking by sentences
        sentences = text.replace('\n', ' ').split('. ')
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_length = len(sentence)
            
            if current_length + sentence_length > chunk_size and current_chunk:
                # Save current chunk
                chunks.append('. '.join(current_chunk) + '.')
                
                # Keep overlap sentences
                overlap_tokens = int(overlap / (chunk_size / len(current_chunk)))
                current_chunk = current_chunk[-overlap_tokens:] if overlap_tokens > 0 else []
                current_length = sum(len(s) for s in current_chunk)
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        # Add remaining chunk
        if current_chunk:
            chunks.append('. '.join(current_chunk) + '.')
        
        return chunks
    
    async def delete_document(self, document_name: str) -> int:
        """
        Delete all chunks of a document.
        
        Args:
            document_name: Name of the document to delete
            
        Returns:
            Number of chunks deleted
        """
        result = await self.db.execute(
            select(PolicyDocument).where(
                PolicyDocument.document_name == document_name
            )
        )
        docs = result.scalars().all()
        
        count = len(docs)
        for doc in docs:
            await self.db.delete(doc)
        
        await self.db.commit()
        
        logger.info("Document deleted from RAG", 
                   document_name=document_name, 
                   chunks_deleted=count)
        
        return count


class DocumentProcessor:
    """Process various document formats for RAG ingestion."""
    
    @staticmethod
    async def process_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    
    @staticmethod
    async def process_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        return text.strip()
    
    @staticmethod
    async def process_txt(file_path: str) -> str:
        """Read a text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    
    @classmethod
    async def process_file(cls, file_path: str) -> str:
        """
        Process a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        file_path_lower = file_path.lower()
        
        if file_path_lower.endswith('.pdf'):
            return await cls.process_pdf(file_path)
        elif file_path_lower.endswith('.docx'):
            return await cls.process_docx(file_path)
        elif file_path_lower.endswith('.txt'):
            return await cls.process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path}")
